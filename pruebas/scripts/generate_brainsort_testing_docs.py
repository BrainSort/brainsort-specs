"""Genera la documentacion de pruebas U4-EJ26 para BrainSort.

Los archivos generados siguen las plantillas entregadas por el docente:
- Plan de pruebas: estructura del PDF "Template-plan-de-pruebas-de-sw".
- Diseno de casos: incluye catalogo de CP y hoja de clases de equivalencia/valores limite.
- Informe: resultados reales de la rama feature/unit-tests-qa.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "pruebas"
EXAMPLE_DIR = OUT_DIR / "ejemplos"

PROJECT = "BrainSort"
VERSION = "1.0"
DATE = "17/05/2026"
AUTHORS = "Equipo BrainSort"

COLORS = {
    "navy": "1F3A5F",
    "blue": "2275C8",
    "light_blue": "D9EAFB",
    "green": "D4EDDA",
    "red": "FFD7D7",
    "yellow": "FFF3CD",
    "gray": "E9ECEF",
    "white": "FFFFFF",
}

TEST_EVIDENCE = {
    "api_unit": "11 suites / 93 tests pasados",
    "api_e2e": "6 suites / 28 tests pasados",
    "app_unit": "6 suites / 31 tests pasados",
    "api_build": "npm run build OK",
    "app_typecheck": "npm run typecheck OK",
    "api_lint": "npm run lint: 0 errores, warnings no bloqueantes",
    "app_lint": "npm run lint: 0 errores, warnings no bloqueantes",
    "api_commit": "fecbfca test: strengthen service coverage review",
    "app_commit": "077d316 fix: align simulation canvas color token",
    "branch": "feature/unit-tests-qa",
}

AUTOMATED_FILES = [
    ("API", "src/auth/auth.service.spec.ts", "Unitaria", 13, "Registro, login usuario/admin, refresh, errores"),
    ("API", "src/users/users.service.spec.ts", "Unitaria", 8, "Perfil y actualizacion"),
    ("API", "src/algorithms/algorithms.service.spec.ts", "Unitaria", 9, "Biblioteca, filtros, detalle"),
    ("API", "src/simulations/simulations.service.spec.ts", "Unitaria", 14, "Simulacion, validacion, engine, errores"),
    ("API", "src/exercises/exercises.service.spec.ts", "Unitaria", 15, "Respuestas, puntos, racha, nivel"),
    ("API", "src/progress/progress.service.spec.ts", "Unitaria", 6, "Progreso, ranking, limites"),
    ("API", "src/badges/badges.service.spec.ts", "Unitaria", 8, "Insignias y criterios"),
    ("API", "src/sync/sync.service.spec.ts", "Unitaria", 8, "Sincronizacion offline"),
    ("API", "src/diagnostics/diagnostics.service.spec.ts", "Unitaria", 5, "Diagnostico inicial"),
    ("API", "src/learning-path/learning-path.service.spec.ts", "Unitaria", 3, "Ruta de aprendizaje"),
    ("API", "src/offline/offline.service.spec.ts", "Unitaria", 4, "Modulos offline"),
    ("API", "test/auth.e2e-spec.ts", "E2E", 9, "Flujo HTTP de autenticacion"),
    ("API", "test/algorithms.e2e-spec.ts", "E2E", 5, "Biblioteca y detalle"),
    ("API", "test/simulations.e2e-spec.ts", "E2E", 5, "Crear simulacion"),
    ("API", "test/exercises.e2e-spec.ts", "E2E", 4, "Ejercicios y respuesta"),
    ("API", "test/progress.e2e-spec.ts", "E2E", 4, "Progreso y ranking"),
    ("API", "test/app.e2e-spec.ts", "E2E", 1, "Health/ruta raiz"),
    ("APP", "packages/core/src/engines/__tests__/bubble-sort.test.ts", "Unitaria", 5, "Bubble Sort"),
    ("APP", "packages/core/src/engines/__tests__/selection-sort.test.ts", "Unitaria", 5, "Selection Sort"),
    ("APP", "packages/core/src/engines/__tests__/insertion-sort.test.ts", "Unitaria", 5, "Insertion Sort"),
    ("APP", "packages/core/src/engines/__tests__/merge-sort.test.ts", "Unitaria", 5, "Merge Sort"),
    ("APP", "packages/core/src/math/__tests__/math.test.ts", "Unitaria", 4, "Escalas, coordenadas, transiciones"),
    ("APP", "packages/core/src/validators/__tests__/dataset.test.ts", "Unitaria", 7, "Validacion de datasets"),
]

CASES = [
    ("CP-AUTH-001", "Auth", "Unitaria", "RF-AUTH", "Registrar usuario valido", "Correo nuevo y rol valido", "nombre/correo/rol/contrasena", "Se crea Usuario y ProgresoUsuario; retorna token, refreshToken y usuario.tipo=usuario", "PASO", "Critica", "src/auth/auth.service.spec.ts"),
    ("CP-AUTH-002", "Auth", "Unitaria", "RF-AUTH", "Rechazar correo duplicado", "Correo ya existe", "correo repetido", "409 ConflictException", "PASO", "Alta", "src/auth/auth.service.spec.ts"),
    ("CP-AUTH-003", "Auth", "Unitaria", "RF-AUTH", "Login usuario valido", "Usuario registrado", "correo/contrasena correctos", "Retorna tokens y tipo=usuario", "PASO", "Critica", "src/auth/auth.service.spec.ts"),
    ("CP-AUTH-004", "Auth", "Unitaria", "RF-AUTH", "Login administrador valido", "Administrador en tabla administradores", "admin@brainsort.edu", "Retorna rol Administrador, tipo=administrador y actualiza ultimoAcceso", "PASO", "Critica", "src/auth/auth.service.spec.ts"),
    ("CP-AUTH-005", "Auth", "E2E", "RF-AUTH", "Flujo register/login/refresh HTTP", "API levantada con Fastify inject", "payloads validos e invalidos", "Endpoints /api/auth responden 201/200/401/409 segun corresponda", "PASO", "Critica", "test/auth.e2e-spec.ts"),
    ("CP-USR-001", "Usuarios", "Unitaria", "RF-USR", "Consultar perfil", "JWT valido", "usuarioId", "Retorna id, nombre, correo, rol", "PASO", "Alta", "src/users/users.service.spec.ts"),
    ("CP-USR-002", "Usuarios", "Unitaria", "RF-USR", "Actualizar perfil", "Usuario existente", "nombre actualizado", "Persiste cambios permitidos", "PASO", "Media", "src/users/users.service.spec.ts"),
    ("CP-ALG-001", "Biblioteca", "Unitaria", "HU-01/CO1", "Listar biblioteca", "Algoritmos activos", "sin filtros", "Retorna categorias, totalAlgoritmos y tarjetas con descripcion <=140", "PASO", "Critica", "src/algorithms/algorithms.service.spec.ts"),
    ("CP-ALG-002", "Biblioteca", "Unitaria", "HU-01", "Filtrar por categoria", "Algoritmos de ordenamiento", "categoria=Ordenamiento", "Consulta Prisma con categoria y activo=true", "PASO", "Alta", "src/algorithms/algorithms.service.spec.ts"),
    ("CP-ALG-003", "Biblioteca", "Unitaria", "HU-01", "Filtrar por nombre", "Algoritmos activos", "search=Bubble", "Busqueda contains insensitive", "PASO", "Alta", "src/algorithms/algorithms.service.spec.ts"),
    ("CP-ALG-004", "Biblioteca", "Unitaria", "HU-01", "Filtrar por tags", "Tags basico, ordenamiento", "tags=basico, ordenamiento", "Usa hasSome con tags recortados", "PASO", "Media", "src/algorithms/algorithms.service.spec.ts"),
    ("CP-ALG-005", "Biblioteca", "E2E", "HU-01/CO1", "GET /api/biblioteca", "API en modo test", "query params validos", "Respuesta 200 con catalogo", "PASO", "Critica", "test/algorithms.e2e-spec.ts"),
    ("CP-ALG-006", "Algoritmos", "Unitaria", "HU-02/CO2", "Detalle de algoritmo", "ID existente", "algoritmoId", "Retorna descripcion completa, complejidad, categoria, tags y pseudocode", "PASO", "Critica", "src/algorithms/algorithms.service.spec.ts"),
    ("CP-ALG-007", "Algoritmos", "Unitaria", "HU-02", "Detalle inexistente", "ID no existe", "missing-id", "404 NotFoundException", "PASO", "Alta", "src/algorithms/algorithms.service.spec.ts"),
    ("CP-SIM-001", "Simulaciones", "Unitaria", "HU-03/CO3", "Crear simulacion personalizada", "Usuario autenticado y algoritmo activo", "valores [5,2,8,1], tamano=4", "Crea SesionSimulacion y retorna pasos/pseudocode/totalPasos", "PASO", "Critica", "src/simulations/simulations.service.spec.ts"),
    ("CP-SIM-002", "Simulaciones", "Unitaria", "HU-03", "Generar datos predeterminados", "tipoOrigen=Predeterminado", "valores vacios o ignorados", "Genera 8 a 15 enteros no preordenados", "PASO", "Alta", "src/simulations/simulations.service.spec.ts"),
    ("CP-SIM-003", "Simulaciones", "Unitaria", "HU-03", "Rechazar tamano inconsistente", "tamano no coincide", "valores length 4, tamano 5", "400 BadRequestException", "PASO", "Alta", "src/simulations/simulations.service.spec.ts"),
    ("CP-SIM-004", "Simulaciones", "Unitaria", "HU-03", "Rechazar valores invalidos", "valores no enteros/fuera de rango", "1.5, null, >999", "400 BadRequestException", "PASO", "Alta", "src/simulations/simulations.service.spec.ts"),
    ("CP-SIM-005", "Simulaciones", "E2E", "HU-04/CO3", "POST /api/simulaciones", "Token valido", "algoritmoId y dataset", "Respuesta 201/200 con pasos validos y estadoActual=Pausa", "PASO", "Critica", "test/simulations.e2e-spec.ts"),
    ("CP-ENG-001", "Engines", "Unitaria", "HU-04", "Bubble Sort ordena y reporta pasos", "Engine disponible", "[5,2,8,1]", "Ultimo estadoArray ordenado y pasos con tipoOperacion", "PASO", "Critica", "packages/core/src/engines/__tests__/bubble-sort.test.ts"),
    ("CP-ENG-002", "Engines", "Unitaria", "HU-04", "Selection Sort ordena", "Engine disponible", "[5,2,8,1]", "Ultimo estadoArray ordenado", "PASO", "Critica", "packages/core/src/engines/__tests__/selection-sort.test.ts"),
    ("CP-ENG-003", "Engines", "Unitaria", "HU-04", "Insertion Sort ordena", "Engine disponible", "[5,2,8,1]", "Ultimo estadoArray ordenado", "PASO", "Critica", "packages/core/src/engines/__tests__/insertion-sort.test.ts"),
    ("CP-ENG-004", "Engines", "Unitaria", "HU-04", "Merge Sort ordena", "Engine disponible", "[5,2,8,1]", "Ultimo estadoArray ordenado", "PASO", "Critica", "packages/core/src/engines/__tests__/merge-sort.test.ts"),
    ("CP-MATH-001", "Core Math", "Unitaria", "HU-04", "Escalas y coordenadas SVG", "Datos numericos", "dataset de barras", "Calcula posiciones y transiciones sin NaN", "PASO", "Alta", "packages/core/src/math/__tests__/math.test.ts"),
    ("CP-VAL-001", "Validadores", "Unitaria", "HU-03", "Dataset valido", "Valores enteros", "[1,2,3]", "Valida sin errores", "PASO", "Alta", "packages/core/src/validators/__tests__/dataset.test.ts"),
    ("CP-VAL-002", "Validadores", "Unitaria", "HU-03", "Dataset invalido", "Valores nulos/no numericos", "[1,null,'x']", "Retorna errores claros", "PASO", "Alta", "packages/core/src/validators/__tests__/dataset.test.ts"),
    ("CP-EJR-001", "Ejercicios", "Unitaria", "RF-EJR", "Listar ejercicios por algoritmo", "Algoritmo con ejercicios", "algoritmoId", "Retorna ejercicios sin respuesta sensible innecesaria", "PASO", "Alta", "src/exercises/exercises.service.spec.ts"),
    ("CP-EJR-002", "Ejercicios", "Unitaria", "RF-EJR", "Responder correctamente", "Ejercicio existente", "respuesta correcta", "correcto=true, puntosGanados, feedbackPositivo", "PASO", "Critica", "src/exercises/exercises.service.spec.ts"),
    ("CP-EJR-003", "Ejercicios", "Unitaria", "RF-EJR", "Responder incorrectamente", "Ejercicio existente", "respuesta incorrecta", "correcto=false, puntosGanados=0, feedbackNegativo", "PASO", "Alta", "src/exercises/exercises.service.spec.ts"),
    ("CP-EJR-004", "Ejercicios", "Unitaria", "RF-EJR", "Actualizar racha", "ultimaActividad ayer/hoy/>2 dias", "fechas limite", "Incrementa, mantiene o reinicia racha segun regla", "PASO", "Critica", "src/exercises/exercises.service.spec.ts"),
    ("CP-EJR-005", "Ejercicios", "E2E", "RF-EJR", "Flujo ejercicio HTTP", "Token valido", "GET y POST responder", "Respuesta contiene correcto, feedback, puntos y progreso", "PASO", "Critica", "test/exercises.e2e-spec.ts"),
    ("CP-PRG-001", "Progreso", "Unitaria", "RF-PRG", "Consultar progreso", "Progreso existente", "usuarioId", "Retorna puntos, nivel, racha, ranking e insignias", "PASO", "Alta", "src/progress/progress.service.spec.ts"),
    ("CP-PRG-002", "Progreso", "Unitaria", "RF-PRG", "Ranking paginado", "Usuarios con progreso", "limit/offset", "Retorna ranking ordenado y total", "PASO", "Alta", "src/progress/progress.service.spec.ts"),
    ("CP-PRG-003", "Progreso", "E2E", "RF-PRG", "GET progreso y ranking", "Token valido", "/api/progreso/me y /ranking", "Respuestas 200 y filtros defensivos", "PASO", "Alta", "test/progress.e2e-spec.ts"),
    ("CP-INS-001", "Insignias", "Unitaria", "RF-INS", "Listar insignias", "Seed de insignias", "sin parametros", "Retorna insignias disponibles", "PASO", "Media", "src/badges/badges.service.spec.ts"),
    ("CP-INS-002", "Insignias", "Unitaria", "RF-INS", "Otorgar insignia por criterio", "Progreso cumple criterio", "usuarioId", "Crea ProgresoInsignia sin duplicar", "PASO", "Alta", "src/badges/badges.service.spec.ts"),
    ("CP-SYNC-001", "Sync", "Unitaria", "RF-SYNC", "Sincronizar sesiones offline", "Payload offline valido", "sesiones[]", "Registra sesiones y puntosActualizados", "PASO", "Alta", "src/sync/sync.service.spec.ts"),
    ("CP-OFF-001", "Offline", "Unitaria", "RF-OFF", "Listar modulos offline", "Algoritmos activos", "usuarioId", "Retorna algoritmoId,nombre,tamanoKB,version,descargado", "PASO", "Media", "src/offline/offline.service.spec.ts"),
    ("CP-OFF-002", "Offline", "Unitaria", "RF-OFF", "Descargar modulo", "Algoritmo existente", "algoritmoId", "Retorna meta,pseudocode,ejercicios", "PASO", "Alta", "src/offline/offline.service.spec.ts"),
    ("CP-DIAG-001", "Diagnostico", "Unitaria", "RF-DIAG", "Listar preguntas", "Preguntas seed", "sin parametros", "Retorna preguntas y opciones", "PASO", "Media", "src/diagnostics/diagnostics.service.spec.ts"),
    ("CP-DIAG-002", "Diagnostico", "Unitaria", "RF-DIAG", "Evaluar diagnostico", "Respuestas enviadas", "indices seleccionados", "Calcula puntaje, guarda resultado y genera ruta", "PASO", "Alta", "src/diagnostics/diagnostics.service.spec.ts"),
    ("CP-RUTA-001", "Ruta Aprendizaje", "Unitaria", "RF-RUTA", "Consultar ruta personalizada", "Usuario con diagnostico/ruta", "usuarioId", "Retorna algoritmos sugeridos en orden", "PASO", "Media", "src/learning-path/learning-path.service.spec.ts"),
]

EQUIVALENCE_ROWS = [
    ("correo", "Email requerido para registro/login", "Formato email valido", "test@example.com", "vacio; sin @; dominio incompleto", '""; test; test@', 0, 1, "normal", 254, 255, 256),
    ("contrasena", "Password requerido, minimo 8 caracteres", "Longitud >= 8", "Password123!", "longitud < 8; vacio", '123; ""', 7, 8, "12 chars", 63, 64, 65),
    ("rol", "Rol de usuario", "Estudiante, Profesor, Autodidacta", "Estudiante", "otro valor; vacio", "Admin; ''", "N/A", "valor valido", "Profesor", "N/A", "Autodidacta", "Administrador"),
    ("conjuntoDeDatos.valores", "Array de enteros para simulacion", "2 a 15 enteros validos", "[5,2,8,1]", "nulos; no numericos; fuera de rango", "[1,null]; [a]; [1000]", 1, 2, 8, 14, 15, 16),
    ("conjuntoDeDatos.tamano", "Debe coincidir con valores.length", "igual a length", "4", "diferente de length; negativo", "5; -1", "length-1", "length", "normal", "length", "length", "length+1"),
    ("tipoOrigen", "Origen de datos", "Predeterminado o Personalizado", "Personalizado", "otro valor; vacio", "Manual; ''", "N/A", "Predeterminado", "Personalizado", "N/A", "Personalizado", "Otro"),
    ("limit ranking", "Limite de resultados de ranking", "entero positivo dentro de rango", "20", "cero; negativo; texto", "0; -1; abc", 0, 1, 20, 99, 100, 101),
]


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        style.font.bold = True


def add_doc_table(doc: Document, headers: list[str], rows: list[tuple | list]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_plan(path: Path) -> None:
    doc = Document()
    set_doc_styles(doc)

    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PLAN DE PRUEBAS DE SOFTWARE")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Conforme a ISO/IEC/IEEE 29119, ISTQB e IEEE 829-2008").italic = True
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Proyecto: {PROJECT}\nVersion del documento: {VERSION}\nFecha: {DATE}\nAutor(es): {AUTHORS}")
    doc.add_page_break()

    doc.add_heading("1. INTRODUCCION", level=1)
    doc.add_heading("1.1 Proposito del documento", level=2)
    doc.add_paragraph(
        "Este documento define el marco de trabajo, la estrategia, el alcance y los recursos necesarios "
        "para verificar y validar BrainSort durante la construccion del proyecto. El SUT incluye la app "
        "React Native/Expo, la API NestJS/Fastify, el paquete core de algoritmos y la persistencia PostgreSQL/Prisma."
    )
    add_bullets(doc, [
        "Identificar el Sistema Bajo Prueba: brainsort-app, brainsort-api y packages/core.",
        "Definir objetivos de calidad funcionales y no funcionales.",
        "Establecer criterios Go/No-Go para la entrega U4-EJ26.",
        "Alinear los casos de prueba con historias de usuario, contratos CO1/CO2/CO3 y codigo automatizado.",
    ])
    doc.add_heading("1.2 Referencias", level=2)
    add_doc_table(doc, ["Documento", "Uso en el plan"], [
        ("Template-plan-de-pruebas-de-sw (full).docx.pdf", "Estructura del plan de pruebas."),
        ("Ejercicio 1. Tecnicas de prueba, Clases de equivalencia y valores al limite.xlsx", "Tecnicas de equivalencia y valores limite para 3.2."),
        ("constitution.md", "Stack, colores de simulacion, convenciones y criterios de arquitectura."),
        ("features/*.spec.md", "Requisitos funcionales y no funcionales."),
        ("plan-de-implementacion/04-contratos-api.md", "Contratos REST entre app y API."),
        ("brainsort-api y brainsort-app", "Codigo fuente y casos de prueba automatizados."),
    ])

    doc.add_heading("2. CONTEXTO DEL PROYECTO", level=1)
    doc.add_heading("2.1 Resumen del sistema bajo prueba", level=2)
    doc.add_paragraph(
        "BrainSort es una aplicacion educativa para aprender algoritmos mediante biblioteca, simulaciones "
        "paso a paso, ejercicios de prediccion, progreso, insignias, diagnostico inicial, ruta de aprendizaje "
        "y modo offline. El frontend se implementa con React Native/Expo y el backend con NestJS sobre Fastify."
    )
    add_doc_table(doc, ["Componente", "Tecnologia", "Responsabilidad"], [
        ("brainsort-app", "Expo SDK 51, React Native 0.74, TypeScript", "UI, hooks MVVM, simulacion visual y consumo API."),
        ("packages/core", "TypeScript", "Engines Bubble, Selection, Insertion, Merge, Stack, Queue, Linked List y utilidades matematicas."),
        ("brainsort-api", "NestJS 10, Fastify, Prisma", "REST API, autenticacion, simulaciones, ejercicios, progreso y offline."),
        ("Base de datos", "PostgreSQL 15+", "Usuarios, algoritmos, progreso, insignias, sesiones, diagnostico y rutas."),
    ])
    doc.add_heading("2.2 Objetivos de negocio", level=2)
    add_bullets(doc, [
        "Permitir que estudiantes visualicen algoritmos con datos predeterminados o personalizados.",
        "Medir aprendizaje mediante ejercicios, progreso, niveles, ranking e insignias.",
        "Mantener calidad verificable por pruebas unitarias, integracion y sistema.",
        "Evitar entregar dependencias, credenciales o datos confidenciales en la linea base.",
    ])
    doc.add_heading("2.3 Stakeholders", level=2)
    add_doc_table(doc, ["Rol", "Expectativas"], [
        ("Docente evaluador", "Evidencias completas, coherentes y usando plantillas indicadas."),
        ("Equipo BrainSort", "Pruebas automatizadas reproducibles y documentacion alineada al codigo."),
        ("Usuario estudiante", "Flujos estables de registro, biblioteca, simulacion y ejercicios."),
        ("QA / desarrollador", "Trazabilidad entre requisito, caso, archivo automatizado y resultado."),
    ])

    doc.add_heading("3. ESTRATEGIA DE PRUEBAS", level=1)
    doc.add_heading("3.1 Objetivos del testing", level=2)
    add_bullets(doc, [
        "Verificar contratos REST de autenticacion, biblioteca, simulaciones, ejercicios, progreso, offline, diagnostico y ruta.",
        "Validar engines y utilidades core con datasets normales, limite e invalidos.",
        "Confirmar que la app compila por typecheck y que los tests automatizados pasan.",
        "Registrar evidencia de resultados para el informe 3.3.",
    ])
    doc.add_heading("3.2 Niveles de prueba", level=2)
    add_doc_table(doc, ["Nivel", "Responsable", "Foco", "Herramienta"], [
        ("Unitarias", "Developer/QA", "Servicios, engines, math, validators.", "Jest, ts-jest, @nestjs/testing"),
        ("Integracion", "QA / Backend", "Endpoints HTTP con Fastify inject y modulos Nest.", "Jest E2E"),
        ("Sistema", "QA", "Ejecucion local app+API, Swagger y navegador Expo Web.", "npm scripts, navegador"),
        ("Regresion", "Equipo", "Re-ejecucion despues de fixes en rama feature/unit-tests-qa.", "Jest, build, typecheck, lint"),
    ])
    doc.add_heading("3.3 Tipos de prueba", level=2)
    add_bullets(doc, [
        "Funcionales: flujos principales de usuario y API.",
        "Caja negra: entradas validas/invalidas en DTOs y endpoints.",
        "Caja blanca: servicios con mocks de Prisma y dependencias.",
        "Valores limite y clases de equivalencia: correo, contrasena, rol, dataset, tamano, tipoOrigen, paginacion.",
        "Regresion automatizada: ejecucion completa de suites existentes.",
    ])
    doc.add_heading("3.4 Estrategia de automatizacion", level=2)
    add_bullets(doc, [
        "Automatizar logica determinista, reglas de negocio, DTOs, engines y contratos HTTP.",
        "Mantener pruebas en repos de codigo, no en el repositorio de documentacion.",
        "No versionar dependencias descargadas por npm ni credenciales.",
        f"Rama usada para evidencias: {TEST_EVIDENCE['branch']}.",
    ])
    doc.add_heading("3.5 Estrategia de datos de prueba", level=2)
    add_doc_table(doc, ["Dato", "Valores"], [
        ("Usuario", "test@example.com, Password123!, roles Estudiante/Profesor/Autodidacta."),
        ("Dataset base", "[5, 2, 8, 1, 9, 3, 7, 4]."),
        ("Limites dataset", "min 2 elementos, max 15 elementos, enteros validos."),
        ("Invalidos", "correo sin @, password <8, tamano distinto, null/no entero/fuera de rango."),
    ])
    doc.add_heading("3.6 Gestion de defectos", level=2)
    add_bullets(doc, [
        "Cada defecto debe incluir ID, severidad, modulo, caso relacionado, resultado esperado, resultado real y resolucion.",
        "Los defectos corregidos se evidencian con commit y re-ejecucion de pruebas.",
        "Criterio de cierre: re-test PASO y cero defectos criticos abiertos.",
    ])
    doc.add_heading("3.7 Estrategia de ambientes", level=2)
    add_doc_table(doc, ["Ambiente", "Uso"], [
        ("Local Windows 11", "Ejecucion de npm test, build, lint y Expo Web."),
        ("Backend local", "http://localhost:3000, Swagger en /api/docs."),
        ("Frontend local", "http://localhost:8081 con API_URL=http://localhost:3000."),
        ("CI/PR", "Validacion esperada en GitHub Actions para PR hacia dev/main."),
    ])

    doc.add_heading("4. ALCANCE DETALLADO", level=1)
    doc.add_heading("4.1 Dentro del alcance", level=2)
    add_bullets(doc, [
        "Auth, Users, Biblioteca/Algoritmos, Simulaciones, Engines, Ejercicios, Progreso, Insignias, Sync, Offline, Diagnostico y Ruta de Aprendizaje.",
        "Codigo fuente de casos automatizados en brainsort-api y brainsort-app.",
        "Documentacion de pruebas 3.1, 3.2 y 3.3 para entrega U4-EJ26.",
    ])
    doc.add_heading("4.2 Fuera de alcance", level=2)
    add_bullets(doc, [
        "Pruebas de penetracion formales y auditoria OWASP completa.",
        "Pruebas en dispositivos fisicos iOS/Android reales.",
        "Pruebas de carga masiva con miles de usuarios concurrentes.",
        "CRUD administrativo no implementado como endpoint en el codigo actual.",
    ])
    doc.add_heading("4.3 Matriz de trazabilidad", level=2)
    add_doc_table(doc, ["Requisito", "Nombre", "Casos"], [
        ("HU-01/CO1", "Navegar biblioteca", "CP-ALG-001 a CP-ALG-005"),
        ("HU-02/CO2", "Detalle de algoritmo", "CP-ALG-006 a CP-ALG-007"),
        ("HU-03", "Datos predeterminados/personalizados", "CP-SIM-001 a CP-SIM-004, CP-VAL-*"),
        ("HU-04/CO3", "Simulacion y engines", "CP-SIM-005, CP-ENG-001 a CP-ENG-004, CP-MATH-001"),
        ("RF-AUTH", "Autenticacion", "CP-AUTH-001 a CP-AUTH-005"),
        ("RF-EJR/RF-PRG", "Ejercicios y progreso", "CP-EJR-*, CP-PRG-*, CP-INS-*"),
        ("RF-OFF/RF-SYNC", "Offline y sincronizacion", "CP-OFF-*, CP-SYNC-*"),
    ])

    doc.add_heading("5. DISENO DE PRUEBAS", level=1)
    doc.add_heading("5.1 Tecnicas utilizadas", level=2)
    add_bullets(doc, [
        "Particiones de equivalencia para valores validos e invalidos de formularios y DTOs.",
        "Analisis de valores limite para longitudes, tamanos de dataset y paginacion.",
        "Pruebas de transicion de estado para simulacion, rachas y progreso.",
        "Pruebas de contrato para endpoints REST y estructura de respuesta.",
    ])
    doc.add_heading("5.2 Priorizacion", level=2)
    add_doc_table(doc, ["Bloque", "Prioridad", "Contenido"], [
        ("1", "Critica", "Auth, biblioteca, simulaciones, engines, ejercicios, build/typecheck."),
        ("2", "Alta", "Progreso, ranking, validadores, offline, sync, diagnostico."),
        ("3", "Media", "Insignias, metricas, casos complementarios y documentacion de evidencias."),
    ])
    doc.add_heading("5.3 Cobertura", level=2)
    add_doc_table(doc, ["Metrica", "Umbral", "Resultado usado"], [
        ("Suites API unitarias", "100% pasadas", TEST_EVIDENCE["api_unit"]),
        ("Suites API E2E", "100% pasadas", TEST_EVIDENCE["api_e2e"]),
        ("Suites frontend core", "100% pasadas", TEST_EVIDENCE["app_unit"]),
        ("Build/typecheck", "Sin errores", f"{TEST_EVIDENCE['api_build']} / {TEST_EVIDENCE['app_typecheck']}"),
        ("Defectos criticos abiertos", "0", "0"),
    ])

    doc.add_heading("6. ENTORNO DE PRUEBAS", level=1)
    doc.add_heading("6.1 Hardware requerido", level=2)
    add_bullets(doc, ["Equipo de desarrollo Windows 11 con Node.js 20+, npm y navegador moderno.", "Opcional: dispositivo Android/iOS para validacion movil manual."])
    doc.add_heading("6.2 Software requerido", level=2)
    add_doc_table(doc, ["Software", "Version/Uso"], [
        ("Node.js / npm", "Node 20+ / npm 10+"),
        ("Backend", "NestJS 10, Fastify, Prisma 5"),
        ("Frontend", "Expo SDK 51, React Native 0.74, React 18"),
        ("DB", "PostgreSQL 15+"),
        ("Testing", "Jest, ts-jest, @nestjs/testing, Fastify inject"),
    ])
    doc.add_heading("6.3 Herramientas de testing", level=2)
    add_bullets(doc, ["Jest para unitarias y e2e.", "Swagger/OpenAPI para validar contratos.", "ESLint/Prettier para calidad estatica.", "TypeScript typecheck para consistencia de tipos."])

    doc.add_heading("7. PLAN DE EJECUCION", level=1)
    doc.add_heading("7.1 Cronograma", level=2)
    add_doc_table(doc, ["Fase", "Duracion estimada", "Hito"], [
        ("Analisis y diseno", "1 dia", "Plan 3.1 y catalogo inicial 3.2."),
        ("Automatizacion", "2 dias", "Suites unitarias/e2e en API y app."),
        ("Ejecucion ciclo 1", "1 dia", "Resultados 3.3 con evidencias."),
        ("Correccion y regresion", "1 dia", "Commits en feature/unit-tests-qa y re-test."),
        ("Cierre", "1 dia", "Documentacion y zip de linea base."),
    ])
    doc.add_heading("7.2 Entry criteria", level=2)
    add_bullets(doc, ["Repos actualizados desde main.", "Dependencias instaladas localmente pero no versionadas.", "Variables .env configuradas sin exponer secretos reales.", "Plan de pruebas aprobado."])
    doc.add_heading("7.3 Exit criteria", level=2)
    add_bullets(doc, ["100% de casos automatizados ejecutados.", "0 fallos en suites unitarias/e2e.", "0 errores en build/typecheck/lint.", "0 defectos criticos abiertos.", "Documentacion 3.1, 3.2 y 3.3 actualizada."])
    doc.add_heading("7.4 Entregables", level=2)
    add_bullets(doc, ["3.1-Plan-de-Pruebas-BrainSort.docx.", "3.2-Casos-de-Prueba-BrainSort.xlsx.", "3.3-Informe-de-Prueba-BrainSort.xlsx.", "Codigo fuente de pruebas automatizadas en brainsort-api y brainsort-app.", "Zip LB-U4-EJ26-REPO-DOCS-BrainSort.zip y LB-U4-EJ26-REPO-CODE-BrainSort.zip."])

    doc.add_heading("8. METRICAS Y REPORTES", level=1)
    doc.add_heading("8.1 Metricas de calidad", level=2)
    add_doc_table(doc, ["Metrica", "Valor"], [
        ("Total automatizados ejecutados", "152 tests"),
        ("API unitarios", TEST_EVIDENCE["api_unit"]),
        ("API E2E", TEST_EVIDENCE["api_e2e"]),
        ("Frontend core", TEST_EVIDENCE["app_unit"]),
        ("Tasa de exito automatizada", "100%"),
        ("Defectos criticos abiertos", "0"),
    ])
    doc.add_heading("8.2 Metricas de proceso", level=2)
    add_bullets(doc, ["Tiempo de re-test: una ejecucion por repo despues de correcciones.", "Evidencias de cambios: commits fecbfca y 077d316 en feature/unit-tests-qa.", "Defect leakage esperado: 0 para flujos cubiertos por automatizacion."])

    doc.add_heading("9. ORGANIZACION DEL EQUIPO", level=1)
    doc.add_heading("9.1 Roles y responsabilidades", level=2)
    add_doc_table(doc, ["Rol", "Responsabilidad"], [
        ("QA Lead", "Planificar, disenar casos, consolidar informe."),
        ("Developer Backend", "Mantener tests API, e2e, build y contratos."),
        ("Developer Frontend", "Mantener tests core, typecheck y flujos Expo."),
        ("Product Owner/Docente", "Validar que la evidencia cumpla la actividad U4-EJ26."),
    ])
    doc.add_heading("9.2 Estructura del equipo", level=2)
    add_doc_table(doc, ["Rol", "Persona(s)", "Contacto"], [
        ("QA Lead", "Equipo BrainSort", "N/A"),
        ("QA Automation", "Equipo BrainSort", "N/A"),
        ("Analista QA", "Equipo BrainSort", "N/A"),
    ])

    doc.save(path)


def style_sheet(ws, widths: list[int] | None = None) -> None:
    header_fill = PatternFill("solid", fgColor=COLORS["navy"])
    header_font = Font(name="Calibri", bold=True, color=COLORS["white"])
    thin = Side(style="thin", color="A6A6A6")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = border
            cell.font = Font(name="Calibri", size=10)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    if widths:
        for idx, width in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(idx)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def build_cases(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Casos de Prueba"
    ws.append(["ID", "Modulo", "Tipo", "Req/HU", "Descripcion", "Precondiciones", "Datos de entrada", "Resultado esperado", "Estado", "Severidad", "Archivo de codigo"])
    for case in CASES:
        ws.append(case)
    style_sheet(ws, [14, 16, 13, 12, 36, 34, 28, 48, 12, 12, 45])

    ws_eq = wb.create_sheet("Equivalencia y Limites")
    ws_eq.append(["Campo", "Descripcion", "Clases validas", "Ej. valido", "Clases invalidas", "Ej. invalido", "MIN-1", "MIN", "NORMAL", "MAX-1", "MAX", "MAX+1"])
    for row in EQUIVALENCE_ROWS:
        ws_eq.append(row)
    style_sheet(ws_eq, [22, 34, 34, 22, 34, 24, 14, 14, 16, 14, 14, 14])

    ws_tr = wb.create_sheet("Trazabilidad")
    ws_tr.append(["Req/HU", "Modulo", "Casos asociados", "Evidencia"])
    mapping: dict[tuple[str, str], list[str]] = {}
    for case in CASES:
        mapping.setdefault((case[3], case[1]), []).append(case[0])
    for (req, module), ids in sorted(mapping.items()):
        ws_tr.append([req, module, ", ".join(ids), "Automatizado" if any(c[10] for c in CASES if c[0] in ids) else "Manual"])
    style_sheet(ws_tr, [16, 18, 80, 18])

    ws_auto = wb.create_sheet("Automatizacion")
    ws_auto.append(["Repo", "Archivo", "Tipo", "Tests", "Cobertura funcional"])
    for row in AUTOMATED_FILES:
        ws_auto.append(row)
    style_sheet(ws_auto, [10, 65, 14, 10, 45])

    ws_sum = wb.create_sheet("Resumen")
    ws_sum.append(["Metrica", "Valor"])
    summary_rows = [
        ("Total casos catalogados", len(CASES)),
        ("Total tests automatizados", 152),
        ("API unitarios", TEST_EVIDENCE["api_unit"]),
        ("API E2E", TEST_EVIDENCE["api_e2e"]),
        ("Frontend core", TEST_EVIDENCE["app_unit"]),
        ("Rama evidencia", TEST_EVIDENCE["branch"]),
        ("Commit API", TEST_EVIDENCE["api_commit"]),
        ("Commit APP", TEST_EVIDENCE["app_commit"]),
    ]
    for row in summary_rows:
        ws_sum.append(row)
    style_sheet(ws_sum, [32, 80])

    wb.save(path)


def build_report(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Resumen Ejecutivo"
    ws.append(["Metrica", "Resultado"])
    rows = [
        ("Proyecto", PROJECT),
        ("Fecha de ejecucion", DATE),
        ("Ciclo", "Regresion U4-EJ26 / feature/unit-tests-qa"),
        ("Total automatizados", "152"),
        ("Pasaron", "152"),
        ("Fallaron", "0"),
        ("Bloqueados", "0"),
        ("Tasa de exito", "100%"),
        ("Veredicto", "APROBADO PARA ENTREGA DE EVIDENCIAS"),
    ]
    for row in rows:
        ws.append(row)
    style_sheet(ws, [32, 70])

    ws_res = wb.create_sheet("Resultados Detallados")
    ws_res.append(["Suite", "Comando", "Resultado", "Observaciones"])
    for row in [
        ("API unitarias", "npm test -- --runInBand", "PASO", TEST_EVIDENCE["api_unit"]),
        ("API E2E", "npm run test:e2e -- --runInBand", "PASO", TEST_EVIDENCE["api_e2e"]),
        ("API build", "npm run build", "PASO", TEST_EVIDENCE["api_build"]),
        ("API lint", "npm run lint", "PASO", TEST_EVIDENCE["api_lint"]),
        ("APP unitarias core", "npm test -- --runInBand", "PASO", TEST_EVIDENCE["app_unit"]),
        ("APP typecheck", "npm run typecheck", "PASO", TEST_EVIDENCE["app_typecheck"]),
        ("APP lint", "npm run lint", "PASO", TEST_EVIDENCE["app_lint"]),
    ]:
        ws_res.append(row)
    style_sheet(ws_res, [24, 42, 14, 65])

    ws_cov = wb.create_sheet("Cobertura")
    ws_cov.append(["Modulo", "Evidencia", "Estado"])
    for row in [
        ("AuthService", "register/login/refresh; usuario y administrador", "PASO"),
        ("AlgorithmsService", "biblioteca, filtros, detalle, not found", "PASO"),
        ("SimulationsService", "dataset personalizado/predeterminado, validaciones, error engine", "PASO"),
        ("ExercisesService", "respuesta, puntos, racha, nivel", "PASO"),
        ("Progress/Badges/Sync/Offline/Diagnostics/LearningPath", "Suites unitarias dedicadas", "PASO"),
        ("Core engines", "Bubble, Selection, Insertion, Merge", "PASO"),
        ("Core math/validators", "math.test.ts y dataset.test.ts", "PASO"),
    ]:
        ws_cov.append(row)
    style_sheet(ws_cov, [34, 80, 14])

    ws_def = wb.create_sheet("Registro de Defectos")
    ws_def.append(["ID", "Severidad", "Modulo", "Descripcion", "Estado", "Resolucion"])
    for row in [
        ("DEF-001", "Alta", "APP Simulacion", "Uso de Accent[400] inexistente rompia typecheck.", "Cerrado", TEST_EVIDENCE["app_commit"]),
        ("DEF-002", "Media", "API Tests", "Cobertura de bordes insuficiente en servicios de algoritmos, ejercicios y simulaciones.", "Cerrado", TEST_EVIDENCE["api_commit"]),
        ("DEF-003", "Baja", "Lint", "Warnings no bloqueantes de unsafe-any/unused en codigo existente.", "Aceptado", "No bloquea entrega; 0 errores."),
    ]:
        ws_def.append(row)
    style_sheet(ws_def, [12, 12, 24, 70, 14, 55])

    ws_ev = wb.create_sheet("Evidencia Codigo")
    ws_ev.append(["Repo", "Rama", "Commit", "Descripcion"])
    ws_ev.append(["brainsort-api", TEST_EVIDENCE["branch"], TEST_EVIDENCE["api_commit"], "Tests API fortalecidos y verificados."])
    ws_ev.append(["brainsort-app", TEST_EVIDENCE["branch"], TEST_EVIDENCE["app_commit"], "Typecheck corregido y tests core verificados."])
    style_sheet(ws_ev, [18, 28, 55, 70])

    ws_con = wb.create_sheet("Conclusiones")
    ws_con.append(["Conclusiones", "Detalle"])
    for row in [
        ("Resultado general", "La evidencia automatizada cubre pruebas unitarias, integracion/e2e y sistema local."),
        ("Consistencia", "Los documentos 3.1, 3.2 y 3.3 quedan alineados con el codigo actual y la tarea U4-EJ26."),
        ("Recomendacion", "Mantener branch feature/unit-tests-qa para revision y preparar zips sin node_modules ni credenciales."),
        ("Go/No-Go", "GO para documentacion de pruebas y codificacion de pruebas automatizadas."),
    ]:
        ws_con.append(row)
    style_sheet(ws_con, [28, 90])

    wb.save(path)


def generate_all() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    EXAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [
        OUT_DIR / "3.1-Plan-de-Pruebas-BrainSort.docx",
        OUT_DIR / "3.2-Casos-de-Prueba-BrainSort.xlsx",
        OUT_DIR / "3.3-Informe-de-Prueba-BrainSort.xlsx",
        EXAMPLE_DIR / "3.1-Plan-de-Pruebas-EJEMPLO.docx",
        EXAMPLE_DIR / "3.2-Casos-de-Prueba-EJEMPLO.xlsx",
        EXAMPLE_DIR / "3.3-Informe-de-Prueba-EJEMPLO.xlsx",
    ]
    build_plan(outputs[0])
    build_cases(outputs[1])
    build_report(outputs[2])
    build_plan(outputs[3])
    build_cases(outputs[4])
    build_report(outputs[5])
    for output in outputs:
        print(f"OK: {output}")


if __name__ == "__main__":
    generate_all()
