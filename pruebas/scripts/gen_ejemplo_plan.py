"""Genera Plan de Pruebas llenado con datos reales de BrainSort"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11)
for i in range(1,4):
    doc.styles[f'Heading {i}'].font.color.rgb = RGBColor(0x1F,0x3A,0x5F)

def tbl(doc, hdrs, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hdrs):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            cell = t.rows[ri+1].cells[ci]; cell.text = str(v)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph()

# PORTADA
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PLAN DE PRUEBAS DE SOFTWARE'); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(0x1F,0x3A,0x5F)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BrainSort - Sprint 3: Simulacion Visual'); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x2E,0x74,0xB5)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Aplicacion educativa para visualizacion de algoritmos de ordenamiento'); r.font.size = Pt(12)
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Version 1.0  |  Mayo 2026').font.size = Pt(12)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Basado en IEEE 829-2008 y especificaciones brainsort-specs'); r.font.size = Pt(10); r.italic = True
doc.add_page_break()

# 1. INTRODUCCION
doc.add_heading('1. Introduccion', level=1)
doc.add_heading('1.1 Proposito', level=2)
doc.add_paragraph('Este documento define la estrategia de pruebas para el modulo de Simulacion Visual del sistema BrainSort, correspondiente al Sprint 3. Cubre los componentes SimulationScreen, SimulationContext, ControlBar, SpeedSlider, PseudocodePanel, CompletionOverlay y los engines de Bubble Sort, Selection Sort e Insertion Sort.')
doc.add_heading('1.2 Alcance', level=2)
doc.add_paragraph('Frontend (brainsort-app): SimulationScreen.tsx, SimulationContext.tsx, ControlBar.tsx, SpeedSlider.tsx, PseudocodePanel.tsx, CompletionOverlay.tsx, BubbleSortAnimation.tsx, SelectionSortAnimation.tsx, InsertionSortAnimation.tsx', style='List Bullet')
doc.add_paragraph('Engine: mock-bubble-sort.ts (generateBubbleSortSteps), mock-selection-sort.ts, mock-insertion-sort.ts', style='List Bullet')
doc.add_paragraph('NO se probara: Backend API, autenticacion, gamificacion, offline, base de datos', style='List Bullet')

doc.add_heading('1.3 Referencias', level=2)
tbl(doc, ['Documento','Version','Fecha'], [
    ['library-simulation.spec.md','1.2','10/04/2026'],
    ['library-simulation.plan.md','1.1','12/04/2026'],
    ['constitution.md','1.0','15/02/2026'],
    ['BrainSort-Historias_de_Usuario.docx','3.0','01/03/2026'],
    ['BrainSort-Contratos.docx','2.0','01/03/2026'],
])

# 2. ELEMENTOS A PROBAR
doc.add_heading('2. Elementos a Probar', level=1)
tbl(doc, ['ID','Modulo','Componente/Archivo','HU','Prioridad'], [
    ['EP-01','Simulacion Visual','SimulationCanvas.tsx + BubbleSortAnimation.tsx','HU-04','Critica'],
    ['EP-02','Simulacion Visual','SelectionSortAnimation.tsx','HU-04','Critica'],
    ['EP-03','Simulacion Visual','InsertionSortAnimation.tsx','HU-04','Critica'],
    ['EP-04','Controles','ControlBar.tsx (Play/Pausa/Reiniciar)','HU-04, HU-06','Critica'],
    ['EP-05','Controles','SpeedSlider.tsx (0.25x-2.0x)','HU-04','Alta'],
    ['EP-06','Pseudocodigo','PseudocodePanel.tsx','HU-04','Alta'],
    ['EP-07','Datos','Generador arreglos aleatorios (8-15 elem)','HU-03','Alta'],
    ['EP-08','Finalizacion','CompletionOverlay.tsx','HU-07','Alta'],
    ['EP-09','Contexto','SimulationContext.tsx (estado global)','HU-04','Critica'],
    ['EP-10','Engine','mock-bubble-sort.ts (generateBubbleSortSteps)','HU-04','Critica'],
    ['EP-11','Engine','mock-selection-sort.ts','HU-04','Critica'],
    ['EP-12','Engine','mock-insertion-sort.ts','HU-04','Critica'],
    ['EP-13','Rendimiento','FPS durante animacion','HU-04','Critica'],
])

# 3. EXCLUSIONES
doc.add_heading('3. Elementos Excluidos', level=1)
for e in ['Backend API (brainsort-api) - Se probara en Sprint 4',
          'Autenticacion (AuthContext) - Probado en Sprint 2',
          'Gamificacion y ejercicios de prediccion - Sprint 5',
          'Funcionalidad offline - Sprint 6',
          'Algoritmos de busqueda - Fuera de alcance v1']:
    doc.add_paragraph(e, style='List Bullet')

# 4. ENFOQUE
doc.add_heading('4. Enfoque de Pruebas', level=1)
tbl(doc, ['Nivel','Que se prueba','Herramienta','Responsable'], [
    ['Unitarias','generateBubbleSortSteps, generateSelectionSortSteps, generateInsertionSortSteps, setSpeed, nextStep','Jest + ts-jest','Carlos Mendoza'],
    ['Integracion','SimulationContext + Engine + Animaciones','React Testing Library + Jest','Ana Rodriguez'],
    ['Sistema (E2E)','Flujo: seleccionar algo -> cargar datos -> Play -> animar -> completar -> overlay','Manual en Expo Go (Android Pixel 5a)','Luis Perez'],
    ['Rendimiento','FPS durante animacion con 15 elementos a 2.0x','React DevTools Profiler','Carlos Mendoza'],
])

# 5. CRITERIOS
doc.add_heading('5. Criterios de Aceptacion', level=1)
tbl(doc, ['Criterio','Metrica','Umbral'], [
    ['Cobertura engine','Jest --coverage en src/engine/','>=85% lineas'],
    ['Casos criticos','CP-SIM-01,02,03,06,09 + CP-REN-01','100% PASADOS (6/6)'],
    ['Casos totales','Todos los CPs del Sprint 3','>=95% (min 18/19)'],
    ['Defectos bloqueantes','Severidad Critica abiertos','0'],
    ['FPS simulacion','Bubble Sort 15 elem, Pixel 5a, 2.0x','>=24 FPS promedio'],
    ['Tiempo carga SimulationScreen','Desde tap en tarjeta hasta barras visibles','<3 segundos'],
])
doc.add_heading('Criterios de Rechazo', level=2)
for c in ['Mas de 1 defecto Critico sin resolver',
          'Build roto (npm run typecheck falla)',
          'FPS <18 en cualquier algoritmo con 15 elementos']:
    doc.add_paragraph(c, style='List Bullet')

# 6. AMBIENTE
doc.add_heading('6. Ambiente de Pruebas', level=1)
tbl(doc, ['Componente','Especificacion'], [
    ['SO Desarrollo','Windows 11 Pro'],
    ['Runtime','Node.js v20.18.0, npm v10.9.2'],
    ['Frontend','Expo SDK 51, React Native 0.74.5, React 18.2.0'],
    ['Dispositivo movil','Google Pixel 5a (Android 14, 6GB RAM)'],
    ['Navegador web','Chrome 126, Edge 126'],
    ['IDE','VS Code 1.89 + ESLint + Prettier'],
    ['Herramientas prueba','Jest 29.x, React Testing Library, React DevTools Profiler'],
])
doc.add_heading('Datos de Prueba', level=2)
for d in ['Arreglo base: [5, 2, 8, 1, 9, 3, 7, 4]',
          'Arreglo minimo: [3, 1] (2 elementos)',
          'Arreglo maximo: [15,14,13,12,11,10,9,8,7,6,5,4,3,2,1] (15 elem, orden inverso)',
          'Arreglo ya ordenado: [1, 2, 3, 4, 5, 6, 7, 8]',
          'Arreglo con duplicados: [5, 3, 5, 1, 3, 8, 1]',
          'Datos invalidos: [5, abc, , 3], campo vacio, caracteres especiales']:
    doc.add_paragraph(d, style='List Bullet')

# 7. RESPONSABILIDADES
doc.add_heading('7. Responsabilidades', level=1)
tbl(doc, ['Nombre','Rol','Responsabilidad'], [
    ['Luis Perez Garcia','QA Lead','Disenar plan y casos, ejecutar pruebas manuales, reportar'],
    ['Carlos Mendoza','Dev Frontend / Engine','Pruebas unitarias del engine, profiling de rendimiento'],
    ['Ana Rodriguez','Dev Frontend','Pruebas de integracion componentes + contexto'],
    ['Dr. Roberto Sanchez','Tech Lead','Revisar plan, aprobar criterios'],
    ['Dra. Maria Lopez','Product Owner','Validar UAT, aprobar release'],
])

# 8. CRONOGRAMA
doc.add_heading('8. Cronograma', level=1)
tbl(doc, ['Fase','Inicio','Fin','Responsable'], [
    ['Preparacion ambiente','20/05/2026','21/05/2026','Luis Perez'],
    ['Pruebas unitarias engine','22/05/2026','24/05/2026','Carlos Mendoza'],
    ['Pruebas integracion','25/05/2026','27/05/2026','Ana Rodriguez'],
    ['Pruebas manuales E2E','28/05/2026','30/05/2026','Luis Perez'],
    ['Pruebas rendimiento','31/05/2026','31/05/2026','Carlos Mendoza'],
    ['Informe final','01/06/2026','02/06/2026','Luis Perez'],
])

# 9. RIESGOS
doc.add_heading('9. Riesgos y Mitigaciones', level=1)
tbl(doc, ['Riesgo','Prob.','Impacto','Mitigacion'], [
    ['Engine genera bucle infinito con datos corruptos','Media','Critico','Implementar timeout en SimulationContext (max 10000 pasos)'],
    ['FPS <24 en Insertion Sort con 15 elem','Media','Alto','Profiling con React DevTools, reducir re-renders en InsertionSortAnimation.tsx'],
    ['PseudocodePanel se desfasa del paso actual','Baja','Alto','Verificar lineaPseudocodigo en cada step del engine'],
    ['CompletionOverlay no desaparece a los 5s','Baja','Medio','Verificar setTimeout en CompletionOverlay.tsx'],
    ['Datos predeterminados generan arreglo ya ordenado','Baja','Medio','Validacion en generador: shuffle hasta que no este ordenado'],
])

# 10. APROBACIONES
doc.add_heading('10. Aprobaciones', level=1)
tbl(doc, ['Nombre','Rol','Firma','Fecha'], [
    ['Luis Perez Garcia','QA Lead','[firmado]','19/05/2026'],
    ['Dr. Roberto Sanchez','Tech Lead','[firmado]','19/05/2026'],
    ['Dra. Maria Lopez','Product Owner','[pendiente]',''],
])

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'ejemplos', '3.1-Plan-de-Pruebas-EJEMPLO.docx')
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print('OK: ' + out)
